"""Text extraction and chunking helpers for the RAG pipeline.

RAG 流水线使用的文本解析与切片工具。
"""

from __future__ import annotations

from io import BytesIO, StringIO
from pathlib import Path
import csv

from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader

from superset_agent_service.config import settings


TEXT_EXTENSIONS = {".txt", ".md", ".markdown", ".json", ".sql", ".log"}
WORD_EXTENSIONS = {".docx"}
EXCEL_EXTENSIONS = {".xlsx", ".xlsm"}
CSV_EXTENSIONS = {".csv"}
PDF_EXTENSIONS = {".pdf"}


def extract_text(filename: str, data: bytes) -> str:
    """Extract readable text from an uploaded file.

    从上传文件中提取可读文本。
    """

    suffix = Path(filename).suffix.lower()
    if suffix in TEXT_EXTENSIONS:
        return _decode_text(data)
    if suffix in WORD_EXTENSIONS:
        return _extract_docx(data)
    if suffix in EXCEL_EXTENSIONS:
        return _extract_xlsx(data)
    if suffix in CSV_EXTENSIONS:
        return _extract_csv(data)
    if suffix in PDF_EXTENSIONS:
        return _extract_pdf(data)
    raise ValueError(
        f"Unsupported file type {suffix or '<none>'}. "
        "Supported: txt, md, csv, xlsx, docx, pdf."
    )


def chunk_text(text: str) -> list[str]:
    """Split text into overlapping chunks suitable for embedding.

    将文本拆分成适合生成向量的、有重叠窗口的切片。
    """

    normalized = "\n".join(line.strip() for line in text.splitlines())
    normalized = "\n".join(line for line in normalized.splitlines() if line)
    if not normalized:
        return []

    size = max(settings.RAG_CHUNK_SIZE, 200)
    overlap = min(max(settings.RAG_CHUNK_OVERLAP, 0), size // 2)
    chunks: list[str] = []
    start = 0
    while start < len(normalized):
        end = min(start + size, len(normalized))
        chunks.append(normalized[start:end].strip())
        if end >= len(normalized):
            break
        start = end - overlap
    return [chunk for chunk in chunks if chunk]


def _decode_text(data: bytes) -> str:
    """Decode bytes with common encodings used by uploaded documents.

    使用常见编码解码上传文档的字节内容。
    """

    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="ignore")


def _extract_docx(data: bytes) -> str:
    """Extract paragraphs and table cells from a Word document.

    从 Word 文档中提取段落和表格单元格文本。
    """

    document = Document(BytesIO(data))
    parts: list[str] = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(part for part in parts if part.strip())


def _extract_xlsx(data: bytes) -> str:
    """Extract sheet names and visible cell values from an Excel workbook.

    从 Excel 工作簿中提取工作表名称和可见单元格内容。
    """

    workbook = load_workbook(BytesIO(data), read_only=True, data_only=True)
    parts: list[str] = []
    for sheet in workbook.worksheets:
        parts.append(f"Sheet: {sheet.title}")
        for row in sheet.iter_rows(values_only=True):
            values = [str(value) for value in row if value is not None]
            if values:
                parts.append(" | ".join(values))
    workbook.close()
    return "\n".join(parts)


def _extract_csv(data: bytes) -> str:
    """Extract rows from a CSV file while preserving column order.

    从 CSV 文件中提取行内容，并保留列顺序。
    """

    text = _decode_text(data)
    reader = csv.reader(StringIO(text))
    return "\n".join(" | ".join(cell for cell in row) for row in reader)


def _extract_pdf(data: bytes) -> str:
    """Extract page text from a PDF file.

    从 PDF 文件中提取页面文本。
    """

    reader = PdfReader(BytesIO(data))
    parts: list[str] = []
    for index, page in enumerate(reader.pages, start=1):
        page_text = page.extract_text() or ""
        if page_text.strip():
            parts.append(f"Page {index}\n{page_text}")
    return "\n\n".join(parts)
